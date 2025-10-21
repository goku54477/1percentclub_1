from fastapi import FastAPI, APIRouter, HTTPException
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List
import uuid
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")


# Define Models
class StatusCheck(BaseModel):
    model_config = ConfigDict(extra="ignore")  # Ignore MongoDB's _id field
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class StatusCheckCreate(BaseModel):
    client_name: str

class WaitlistEntry(BaseModel):
    firstName: str = Field(..., min_length=1, max_length=100)
    lastName: str = Field(..., min_length=1, max_length=100)
    email: EmailStr
    phone: str = Field(..., min_length=10, max_length=20)

# Add your routes to the router instead of directly to app
@api_router.get("/")
async def root():
    return {"message": "Hello World"}

@api_router.post("/status", response_model=StatusCheck)
async def create_status_check(input: StatusCheckCreate):
    status_dict = input.model_dump()
    status_obj = StatusCheck(**status_dict)
    
    # Convert to dict and serialize datetime to ISO string for MongoDB
    doc = status_obj.model_dump()
    doc['timestamp'] = doc['timestamp'].isoformat()
    
    _ = await db.status_checks.insert_one(doc)
    return status_obj

@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    # Exclude MongoDB's _id field from the query results
    status_checks = await db.status_checks.find({}, {"_id": 0}).to_list(1000)
    
    # Convert ISO string timestamps back to datetime objects
    for check in status_checks:
        if isinstance(check['timestamp'], str):
            check['timestamp'] = datetime.fromisoformat(check['timestamp'])
    
    return status_checks

@api_router.post("/waitlist")
async def add_to_waitlist(entry: WaitlistEntry):
    """
    Save waitlist entry to a Word document
    """
    try:
        # Define the document path
        doc_path = ROOT_DIR / "waitlist_submissions.docx"
        
        # Check if document exists, if not create it with header
        if not doc_path.exists():
            doc = Document()
            
            # Add title
            title = doc.add_heading('1% Waitlist Submissions', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add creation date
            date_para = doc.add_paragraph()
            date_para.add_run(f'Created: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Empty line
            
            # Save the initial document
            doc.save(str(doc_path))
        
        # Open the document
        doc = Document(str(doc_path))
        
        # Add entry heading
        timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
        heading = doc.add_heading(f'Submission - {timestamp}', level=2)
        
        # Add entry details
        details = [
            ('Name', f'{entry.firstName} {entry.lastName}'),
            ('Email', entry.email),
            ('Phone', entry.phone),
            ('Timestamp', timestamp)
        ]
        
        for label, value in details:
            p = doc.add_paragraph()
            p.add_run(f'{label}: ').bold = True
            p.add_run(value)
        
        # Add separator
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()  # Empty line
        
        # Save the document
        doc.save(str(doc_path))
        
        logger.info(f"Waitlist entry saved: {entry.email}")
        
        return {
            "success": True,
            "message": "Successfully added to waitlist",
            "data": {
                "name": f"{entry.firstName} {entry.lastName}",
                "email": entry.email,
                "timestamp": timestamp
            }
        }
        
    except Exception as e:
        logger.error(f"Error saving waitlist entry: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to save waitlist entry: {str(e)}"
        )

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()